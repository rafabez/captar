"""Render a project's sections into a downloadable DOCX or PDF."""
import re
from html import escape
from io import BytesIO

SECTION_LABELS: dict[str, str] = {
    "summary": "Resumo",
    "justification": "Justificativa",
    "objectives": "Objetivos",
    "target_audience": "Público-alvo",
    "accessibility": "Acessibilidade",
    "counterparts": "Contrapartidas",
    "schedule": "Cronograma",
    "team": "Equipe",
    "communication": "Comunicação",
    "budget": "Orçamento",
}
ORDER = list(SECTION_LABELS.keys())

DOCX_MEDIA = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _slug(name: str) -> str:
    s = re.sub(r"[^a-z0-9]+", "-", (name or "projeto").lower()).strip("-")
    return s or "projeto"


def _ordered(content_by_type: dict[str, str], only: list[str] | None) -> list[tuple[str, str]]:
    return [
        (t, content_by_type[t])
        for t in ORDER
        if t in content_by_type and content_by_type[t].strip()
        and (not only or t in only)
    ]


def _meta(project) -> str:
    bits = [project.area, project.city and f"{project.city}/{project.state or ''}".strip("/")]
    return " · ".join(b for b in bits if b)


def build(project, content_by_type: dict[str, str], fmt: str, only: list[str] | None):
    sections = _ordered(content_by_type, only)
    if not sections:
        raise ValueError("Nenhuma seção com conteúdo para exportar.")
    if fmt == "pdf":
        return _pdf(project, sections)
    return _docx(project, sections)


def _docx(project, sections):
    from docx import Document

    doc = Document()
    doc.add_heading(project.name or "Projeto", 0)
    meta = _meta(project)
    if meta:
        doc.add_paragraph(meta)
    for section_type, content in sections:
        doc.add_heading(SECTION_LABELS.get(section_type, section_type), level=1)
        for para in content.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue(), f"{_slug(project.name)}.docx", DOCX_MEDIA


def _pdf(project, sections):
    from weasyprint import HTML

    body = [f"<h1>{escape(project.name or 'Projeto')}</h1>"]
    meta = _meta(project)
    if meta:
        body.append(f'<p class="meta">{escape(meta)}</p>')
    for section_type, content in sections:
        body.append(f"<h2>{escape(SECTION_LABELS.get(section_type, section_type))}</h2>")
        for para in content.split("\n\n"):
            if para.strip():
                body.append(f"<p>{escape(para.strip())}</p>")

    html = f"""<!doctype html><html><head><meta charset="utf-8"><style>
      @page {{ size: A4; margin: 2.2cm; }}
      body {{ font-family: 'DejaVu Serif', serif; color: #211C16; line-height: 1.5; }}
      h1 {{ font-size: 24pt; color: #211C16; margin-bottom: 4pt; }}
      h2 {{ font-size: 14pt; color: #C4553F; margin-top: 18pt; border-bottom: 1px solid #E6E0D6; padding-bottom: 3pt; }}
      p {{ font-size: 11pt; margin: 6pt 0; text-align: justify; }}
      .meta {{ color: #6E655A; font-size: 10pt; }}
    </style></head><body>{''.join(body)}</body></html>"""

    return HTML(string=html).write_pdf(), f"{_slug(project.name)}.pdf", "application/pdf"
